#!/usr/bin/env python3
"""
Word Document Utilities
Core module for local DOCX manipulation
"""

import os
import json
from pathlib import Path
from typing import Optional, List, Dict, Any

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    Document = None


class DOCXAnalyzer:
    """Analyze Word documents"""

    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"DOCX not found: {file_path}")
        self.doc = Document(str(self.file_path))

    def get_metadata(self) -> Dict[str, Any]:
        """Get document metadata"""
        core_props = self.doc.core_properties

        return {
            'file': str(self.file_path),
            'title': core_props.title or 'N/A',
            'author': core_props.author or 'N/A',
            'subject': core_props.subject or 'N/A',
            'created': str(core_props.created) if core_props.created else 'N/A',
            'modified': str(core_props.modified) if core_props.modified else 'N/A',
            'paragraphs': len(self.doc.paragraphs),
            'tables': len(self.doc.tables),
            'sections': len(self.doc.sections)
        }

    def get_structure(self) -> List[Dict[str, Any]]:
        """Get document structure (headings and sections)"""
        structure = []

        for para in self.doc.paragraphs:
            style_name = para.style.name if para.style else "Normal"

            if style_name.startswith('Heading'):
                structure.append({
                    'type': 'heading',
                    'level': style_name,
                    'text': para.text[:100],
                    'word_count': len(para.text.split())
                })

        return structure

    def extract_text(self) -> str:
        """Extract all text from document"""
        texts = []
        for para in self.doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        return "\n\n".join(texts)

    def get_word_count(self) -> Dict[str, Any]:
        """Get detailed word count"""
        total = 0
        by_section = {}
        current_section = "Intro"

        for para in self.doc.paragraphs:
            style_name = para.style.name if para.style else "Normal"

            if style_name.startswith('Heading 1'):
                current_section = para.text[:50] or f"Section {len(by_section) + 1}"
                if current_section not in by_section:
                    by_section[current_section] = 0

            words = len(para.text.split())
            total += words

            if current_section not in by_section:
                by_section[current_section] = 0
            by_section[current_section] += words

        return {
            'total': total,
            'by_section': by_section
        }

    def get_table_info(self) -> List[Dict[str, Any]]:
        """Get info about all tables"""
        tables = []
        for i, table in enumerate(self.doc.tables, 1):
            rows = len(table.rows)
            cols = len(table.columns)

            # Get header row if exists
            header = []
            if rows > 0:
                for cell in table.rows[0].cells:
                    header.append(cell.text[:30])

            tables.append({
                'table': i,
                'rows': rows,
                'columns': cols,
                'header': header
            })

        return tables

    def analyze(self) -> Dict[str, Any]:
        """Full document analysis"""
        metadata = self.get_metadata()
        word_count = self.get_word_count()
        structure = self.get_structure()
        tables = self.get_table_info()

        return {
            **metadata,
            'word_count': word_count['total'],
            'word_count_by_section': word_count['by_section'],
            'structure': structure,
            'tables_info': tables
        }


class DOCXEditor:
    """Edit Word documents"""

    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"DOCX not found: {file_path}")
        self.doc = Document(str(self.file_path))

    def replace_text(
        self,
        find: str,
        replace: str,
        preserve_formatting: bool = True
    ) -> int:
        """
        Find and replace text

        Returns number of replacements made
        """
        count = 0

        for para in self.doc.paragraphs:
            if find in para.text:
                for run in para.runs:
                    if find in run.text:
                        run.text = run.text.replace(find, replace)
                        count += 1

        # Also check tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if find in para.text:
                            for run in para.runs:
                                if find in run.text:
                                    run.text = run.text.replace(find, replace)
                                    count += 1

        return count

    def update_paragraph(
        self,
        para_index: int,
        new_text: str,
        preserve_style: bool = True
    ):
        """Update specific paragraph"""
        if para_index < 0 or para_index >= len(self.doc.paragraphs):
            raise ValueError(f"Invalid paragraph index: {para_index}")

        para = self.doc.paragraphs[para_index]

        if preserve_style and para.runs:
            # Get first run's formatting
            first_run = para.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            font_bold = first_run.font.bold

            # Clear and add new text
            para.clear()
            run = para.add_run(new_text)

            # Apply formatting
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = font_size
            if font_bold is not None:
                run.font.bold = font_bold
        else:
            para.text = new_text

    def update_table_cell(
        self,
        table_index: int,
        row: int,
        col: int,
        value: str
    ):
        """Update specific table cell"""
        if table_index < 0 or table_index >= len(self.doc.tables):
            raise ValueError(f"Invalid table index: {table_index}")

        table = self.doc.tables[table_index]

        if row < 0 or row >= len(table.rows):
            raise ValueError(f"Invalid row: {row}")
        if col < 0 or col >= len(table.columns):
            raise ValueError(f"Invalid column: {col}")

        cell = table.rows[row].cells[col]
        cell.text = value

    def add_paragraph(
        self,
        text: str,
        style: str = 'Normal',
        after_index: Optional[int] = None
    ):
        """Add new paragraph"""
        if after_index is not None:
            # Insert after specific paragraph
            # Note: This is complex in python-docx, simplified here
            para = self.doc.add_paragraph(text, style=style)
        else:
            para = self.doc.add_paragraph(text, style=style)

        return para

    def add_heading(self, text: str, level: int = 1):
        """Add heading"""
        return self.doc.add_heading(text, level=level)

    def save(self, output_path: str) -> str:
        """Save document"""
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output))
        return str(output)


def analyze_docx(file_path: str) -> Dict[str, Any]:
    """Quick analysis function"""
    analyzer = DOCXAnalyzer(file_path)
    return analyzer.analyze()


def extract_text(file_path: str) -> str:
    """Quick text extraction"""
    analyzer = DOCXAnalyzer(file_path)
    return analyzer.extract_text()


def replace_text(
    file_path: str,
    find: str,
    replace: str,
    output_path: str
) -> int:
    """Find and replace text in DOCX"""
    editor = DOCXEditor(file_path)
    count = editor.replace_text(find, replace)
    editor.save(output_path)
    return count


if __name__ == '__main__':
    import sys

    if len(sys.argv) < 2:
        print("Usage: python docx_utils.py <docx_file>")
        sys.exit(1)

    result = analyze_docx(sys.argv[1])
    print(json.dumps(result, indent=2, default=str))
